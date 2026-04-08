#!/usr/bin/env python3
"""
generate_screens_doc.py
Reads ftf_crtt_2023_archive.py and produces a Word document describing
each screen in the Competitive Reaction Time Task (CRTT) game, including
all editable text and configuration variables.

Requirements: pip install python-docx
"""

import re
import ast
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOURCE_FILE = "ftf_crtt_2023_archive.py"
OUTPUT_FILE = "F_CRTT_Screens_Documentation.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def extract_config(source: str) -> dict:
    """Extract top-level scalar and list assignments from the config section."""
    config = {}
    scalar_patterns = [
        (r'^NUM_GAMES\s*=\s*(\d+)', 'NUM_GAMES', int),
        (r'^NUM_ROUNDS\s*=\s*(\d+)', 'NUM_ROUNDS', int),
        (r'^MAX_WAIT\s*=\s*(\d+)', 'MAX_WAIT', int),
        (r'^MIN_WAIT\s*=\s*(\d+)', 'MIN_WAIT', int),
        (r'^ELAPSED_TIME\s*=\s*(\d+)', 'ELAPSED_TIME', int),
        (r'^RESPONSE_TIMEOUT\s*=\s*(\d+)', 'RESPONSE_TIMEOUT', int),
        (r'^KEY1\s*=\s*"([^"]*)"', 'KEY1', str),
        (r'^KEY2\s*=\s*"([^"]*)"', 'KEY2', str),
        (r'^WINDOW_WIDTH\s*=\s*(\d+)', 'WINDOW_WIDTH', int),
        (r'^WINDOW_HEIGHT\s*=\s*(\d+)', 'WINDOW_HEIGHT', int),
        (r'^FONT_TYPE\s*=\s*"([^"]*)"', 'FONT_TYPE', str),
        (r'^FONT_SIZE\s*=\s*(\d+)', 'FONT_SIZE', int),
        (r'^FONT_COLOUR\s*=\s*"([^"]*)"', 'FONT_COLOUR', str),
        (r'^GO_COLOUR\s*=\s*"([^"]*)"', 'GO_COLOUR', str),
        (r'^BG_COLOUR\s*=\s*"([^"]*)"', 'BG_COLOUR', str),
    ]
    for pattern, name, cast in scalar_patterns:
        m = re.search(pattern, source, re.MULTILINE)
        if m:
            try:
                config[name] = cast(m.group(1))
            except Exception:
                config[name] = m.group(1)

    # Extract DEFAULT_GAME_CONDITIONS list
    m = re.search(r'^DEFAULT_GAME_CONDITIONS\s*=\s*(\[.*?\])', source, re.MULTILINE | re.DOTALL)
    if m:
        try:
            config['DEFAULT_GAME_CONDITIONS'] = ast.literal_eval(m.group(1))
        except Exception:
            pass

    return config


def extract_text_strings(source: str) -> dict:
    """Extract text_* string variables, handling backslash-continued multi-line strings."""
    texts = {}
    # Match:  text_xxx = "..." optionally continued with \ + newline + "..."
    pattern = r'^(text_\w+)\s*=\s*((?:"[^"]*"|\s*\\?\s*)+)'
    for m in re.finditer(pattern, source, re.MULTILINE):
        name = m.group(1)
        raw = m.group(2)
        raw = re.sub(r'\\\s*\n\s*', '', raw)  # join backslash-continued lines
        try:
            value = ast.literal_eval(raw.strip())
            texts[name] = value
        except Exception:
            texts[name] = raw.strip()
    return texts


def set_cell_bg(cell, hex_color: str):
    """Set a table cell's background shading (hex_color e.g. 'DDDDDD')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_screenshot_placeholder(doc):
    """Insert a bordered, shaded table row as a screenshot placeholder."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'CCCCCC')
    cell.width = Inches(6)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[ SCREENSHOT — insert manually ]')
    run.bold = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(13)
    # give the row some height via paragraph spacing
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(28)
    doc.add_paragraph()  # spacer after placeholder


def bold_label(para, label: str, value: str):
    para.add_run(label).bold = True
    para.add_run(value)


def add_var_table(doc, rows):
    """Add a 3-column table: Variable | Default Value | Description."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for cell, heading_text in zip(hdr, ("Variable", "Default Value", "Description")):
        cell.text = heading_text
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for name, val, desc in rows:
        row = tbl.add_row().cells
        row[0].text = str(name)
        row[1].text = str(val) if val is not None else '—'
        row[2].text = desc
    doc.add_paragraph()


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    with open(SOURCE_FILE, 'r') as f:
        source = f.read()

    cfg = extract_config(source)
    txt = extract_text_strings(source)

    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────────
    doc.add_heading('FTF-CRTT Screen Documentation', 0)
    doc.add_paragraph(
        'This document describes every screen in the Competitive Reaction Time Task (CRTT) '
        'game, including all editable copy variables and configuration settings. '
        'Source file: ftf_crtt_2023_archive.py'
    )

    # ── Section 1: Configuration Variables ────────────────────────────────────
    doc.add_heading('1. Configuration Variables', level=1)
    doc.add_paragraph(
        "Defined at the top of the source file under 'Experiment Configuration'. "
        "Change these to alter game behaviour without touching any logic."
    )

    add_var_table(doc, [
        ('NUM_GAMES', cfg.get('NUM_GAMES'),
         'Number of games to play per session'),
        ('DEFAULT_GAME_CONDITIONS', cfg.get('DEFAULT_GAME_CONDITIONS'),
         "Default break condition pre-selected for each game slot. "
         "Must match condition_options values: 'No Break', '5 Seconds', '10 Seconds', '15 Seconds'"),
        ('NUM_ROUNDS', cfg.get('NUM_ROUNDS'),
         'Number of rounds per game'),
        ('MAX_WAIT', cfg.get('MAX_WAIT'),
         'Maximum random wait time before the GO!! signal (seconds)'),
        ('MIN_WAIT', cfg.get('MIN_WAIT'),
         'Minimum random wait time before the GO!! signal (seconds)'),
        ('ELAPSED_TIME', cfg.get('ELAPSED_TIME'),
         'Reaction-time gap threshold (ms). If the second player presses within this '
         'window of the first, the winner is randomised; beyond it the faster player wins.'),
        ('RESPONSE_TIMEOUT', cfg.get('RESPONSE_TIMEOUT'),
         'Seconds to wait for the second player to respond before auto-awarding the win'),
        ('KEY1', cfg.get('KEY1'),
         'Keyboard key assigned to Player 1'),
        ('KEY2', cfg.get('KEY2'),
         'Keyboard key assigned to Player 2'),
        ('WINDOW_WIDTH', cfg.get('WINDOW_WIDTH'),
         'Window width in pixels (used when fullscreen is disabled)'),
        ('WINDOW_HEIGHT', cfg.get('WINDOW_HEIGHT'),
         'Window height in pixels (used when fullscreen is disabled)'),
        ('FONT_TYPE', cfg.get('FONT_TYPE'),
         'Font face used throughout the game'),
        ('FONT_SIZE', cfg.get('FONT_SIZE'),
         'Base font size (points)'),
        ('FONT_COLOUR', cfg.get('FONT_COLOUR'),
         'Default text colour'),
        ('GO_COLOUR', cfg.get('GO_COLOUR'),
         'Background colour when the GO!! signal fires'),
        ('BG_COLOUR', cfg.get('BG_COLOUR'),
         'Default background colour'),
    ])

    # ── Section 2: Editable Text Strings ──────────────────────────────────────
    doc.add_heading('2. Editable Text Strings', level=1)
    doc.add_paragraph(
        "All user-facing copy is defined in the 'Text Strings' section. "
        "Edit these to change what players read without modifying any logic."
    )
    doc.add_paragraph(
        "Strings containing '{}' are format strings — placeholders are replaced at "
        "runtime with player names, counts, or times. Do not remove '{}' unless you "
        "also update the matching .format() call in the code."
    )

    add_var_table(doc, [
        ('text_welcome', txt.get('text_welcome'),
         'Large heading on the Welcome screen'),
        ('text_space_continue', txt.get('text_space_continue'),
         'Instruction shown wherever Space advances the game'),
        ('text_p1_name', txt.get('text_p1_name'),
         'Prompt asking Player 1 to enter their name'),
        ('text_p2_name', txt.get('text_p2_name'),
         'Prompt asking Player 2 to enter their name'),
        ('text_keys', txt.get('text_keys'),
         'Confirms key assignments after name entry. '
         '{0}=P1 name, {1}=P1 key, {2}=P2 name, {3}=P2 key '
         '(defined but not currently displayed in the screen flow)'),
        ('text_instructions', txt.get('text_instructions'),
         'Game rules shown before the first game'),
        ('text_round', txt.get('text_round'),
         'Round label template; {} = round number '
         '(defined but displayed inline in the round header instead)'),
        ('text_get_ready', txt.get('text_get_ready'),
         "Text shown during the 'ready' phase of each round"),
        ('text_get_set', txt.get('text_get_set'),
         "Text shown ~1 second before the GO signal"),
        ('text_go', txt.get('text_go'),
         'The GO signal text (screen also flashes GO_COLOUR)'),
        ('text_win', txt.get('text_win'),
         'Win prompt when there is no forced break. '
         '{0}=winner, {1}=winner (selects blast), {2}=loser'),
        ('text_blast', txt.get('text_blast'),
         'Standby message — defined in source but not currently shown on screen. '
         'Could be displayed during blast delivery if desired.'),
        ('text_win_wait', txt.get('text_win_wait'),
         'Win screen when a forced break is active. '
         '{0}=winner, {1}=winner, {2}=break seconds, {3}=loser'),
        ('text_valid', txt.get('text_valid'),
         'Error shown when an invalid blast level is entered. {0}=loser (on standby)'),
        ('text_game_complete', txt.get('text_game_complete'),
         'Between-game message template (defined but used inline). '
         '{0}=completed game number, {1}=next game number'),
        ('text_game_over', txt.get('text_game_over'),
         'Title on the final summary screen after all games complete'),
    ])

    # ── Section 3: Screen-by-Screen Flow ──────────────────────────────────────
    doc.add_heading('3. Screen-by-Screen Flow', level=1)
    doc.add_paragraph(
        'The game progresses through the following screens in order. '
        'Each entry includes the text displayed, relevant editable variables, '
        'player interaction required, and a placeholder for a screenshot.'
    )

    screens = [
        {
            'title': 'Screen 1 — Experiment Settings',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                'The first thing shown when the program launches. The experimenter '
                'configures the session before any players interact. All values can be '
                'adjusted via spinboxes and dropdowns.'
            ),
            'display': (
                'Experiment Settings  [heading]\n\n'
                'Number of Games:           [spinbox  1–10,  default: NUM_GAMES]\n'
                'Game 1 Break:              [dropdown: No Break | 5 Seconds | 10 Seconds | 15 Seconds]\n'
                'Game 2 Break:              [dropdown] … (one row per game)\n'
                'Rounds per Game:           [spinbox  1–100, default: NUM_ROUNDS]\n'
                'Max Wait (s):              [spinbox  0–60,  default: MAX_WAIT]\n'
                'Min Wait (s):              [spinbox  0–60,  default: MIN_WAIT]\n'
                'No Response Timeout (s):   [spinbox  1–30,  default: RESPONSE_TIMEOUT]\n\n'
                f'{txt.get("text_space_continue", "Press Space to continue")}'
            ),
            'variables': [
                ('NUM_GAMES', cfg.get('NUM_GAMES'), 'Default value in the Number of Games spinbox'),
                ('DEFAULT_GAME_CONDITIONS', cfg.get('DEFAULT_GAME_CONDITIONS'),
                 'Pre-selected break condition for each game slot'),
                ('NUM_ROUNDS', cfg.get('NUM_ROUNDS'), 'Default value in the Rounds per Game spinbox'),
                ('MAX_WAIT', cfg.get('MAX_WAIT'), 'Default value in the Max Wait spinbox'),
                ('MIN_WAIT', cfg.get('MIN_WAIT'), 'Default value in the Min Wait spinbox'),
                ('RESPONSE_TIMEOUT', cfg.get('RESPONSE_TIMEOUT'),
                 'Default value in the No Response Timeout spinbox'),
                ('text_space_continue', txt.get('text_space_continue'),
                 'Label at the bottom of the settings panel'),
            ],
            'interaction': 'Press Space (or Space while a spinbox is focused) to confirm and start.',
        },
        {
            'title': 'Screen 2 — Welcome',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': 'Shown immediately after settings are confirmed.',
            'display': (
                f"{txt.get('text_welcome', 'Welcome!')}"
                f"\n\n{txt.get('text_space_continue', 'Press Space to continue')}"
            ),
            'variables': [
                ('text_welcome', txt.get('text_welcome'), 'Large welcome heading'),
                ('text_space_continue', txt.get('text_space_continue'), 'Instruction to advance'),
            ],
            'interaction': 'Press Space to advance to Player 1 name entry.',
        },
        {
            'title': 'Screen 3 — Player 1 Name Entry',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': 'Prompts Player 1 to type their name into a text field and press Enter.',
            'display': txt.get('text_p1_name', ''),
            'variables': [
                ('text_p1_name', txt.get('text_p1_name'), 'Label above the text entry field'),
                ('KEY1', cfg.get('KEY1'), "Player 1's reaction key (shown for reference)"),
            ],
            'interaction': 'Player 1 types a name and presses Enter.',
        },
        {
            'title': 'Screen 4 — Player 2 Name Entry',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                'Prompts Player 2 to type their name and press Enter. '
                'A timestamped CSV output filename is generated at this point.'
            ),
            'display': txt.get('text_p2_name', ''),
            'variables': [
                ('text_p2_name', txt.get('text_p2_name'), 'Label above the text entry field'),
                ('KEY2', cfg.get('KEY2'), "Player 2's reaction key (shown for reference)"),
            ],
            'interaction': 'Player 2 types a name and presses Enter.',
        },
        {
            'title': 'Screen 5 — Game Instructions (first game only)',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                'Shown before Game 1 begins. Displays the game counter, task instructions, '
                'and the prompt to start.'
            ),
            'display': (
                f"Game 1 of [NUM_GAMES]\n\n"
                f"{txt.get('text_instructions', '')}\n\n"
                f"{txt.get('text_space_continue', '')}"
            ),
            'variables': [
                ('text_instructions', txt.get('text_instructions'),
                 'Rules / instructions text shown to players'),
                ('text_space_continue', txt.get('text_space_continue'), 'Prompt to begin'),
                ('NUM_GAMES', cfg.get('NUM_GAMES'), 'Total game count shown in the header line'),
            ],
            'interaction': 'Press Space to start the first round of Game 1.',
        },
        {
            'title': 'Screen 6 — Between-Game Transition (games 2+)',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                'Shown after each game completes, except after the final game. '
                'Announces the next game number.'
            ),
            'display': (
                'Game [N] complete!\n\n'
                f'Game [N+1] of [NUM_GAMES]\n\n'
                f"{txt.get('text_space_continue', 'Press Space to continue')}"
            ),
            'variables': [
                ('text_space_continue', txt.get('text_space_continue'), 'Prompt to begin next game'),
                ('NUM_GAMES', cfg.get('NUM_GAMES'), 'Total game count shown in the header'),
            ],
            'interaction': 'Press Space to start the next game.',
        },
        {
            'title': 'Screen 7 — Round Ready',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                "Shown at the start of each round. Displays game and round counters, "
                "then the 'Ready…' prompt."
            ),
            'display': (
                f"Game [current game] of [NUM_GAMES]  |  Round [current round] of [NUM_ROUNDS]\n\n"
                f"{txt.get('text_get_ready', 'Ready...')}"
            ),
            'variables': [
                ('text_get_ready', txt.get('text_get_ready'), "The 'ready' prompt copy"),
                ('NUM_GAMES', cfg.get('NUM_GAMES'), 'Shown in round header'),
                ('NUM_ROUNDS', cfg.get('NUM_ROUNDS'), 'Shown in round header'),
            ],
            'interaction': "No player input — advances automatically to 'Set…' after ~1 second.",
        },
        {
            'title': 'Screen 8 — Get Set',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                "Shown approximately 1 second after the Ready screen, just before the GO signal fires."
            ),
            'display': (
                f"Game [current game] of [NUM_GAMES]  |  Round [current round] of [NUM_ROUNDS]\n\n"
                f"{txt.get('text_get_set', 'Set...')}"
            ),
            'variables': [
                ('text_get_set', txt.get('text_get_set'), "The 'set' prompt copy"),
                ('MIN_WAIT', cfg.get('MIN_WAIT'),
                 'Minimum additional random wait before GO fires (seconds)'),
                ('MAX_WAIT', cfg.get('MAX_WAIT'),
                 'Maximum additional random wait before GO fires (seconds)'),
            ],
            'interaction': (
                f"No player input — GO fires after a random delay between "
                f"MIN_WAIT ({cfg.get('MIN_WAIT')}s) and MAX_WAIT ({cfg.get('MAX_WAIT')}s)."
            ),
        },
        {
            'title': 'Screen 9 — GO!! (Reaction Phase)',
            'bg': f"Orange / GO_COLOUR ({cfg.get('GO_COLOUR')})",
            'description': (
                'The screen background flashes to GO_COLOUR and the display text is cleared. '
                'Both players must press their assigned key as fast as possible. '
                'The first player to press records their timestamp; the game then waits up to '
                f'RESPONSE_TIMEOUT ({cfg.get("RESPONSE_TIMEOUT")}s) for the second player. '
                'If no second press arrives within the timeout, the first presser wins automatically.'
            ),
            'display': (
                f'[Screen flashes {cfg.get("GO_COLOUR", "orange")} — display text is cleared]\n\n'
                f'(text_go = "{txt.get("text_go", "GO!!")}" is defined but the text area is '
                f'emptied when the flash fires; text_go could be added here if desired)'
            ),
            'variables': [
                ('GO_COLOUR', cfg.get('GO_COLOUR'), 'Flash background colour'),
                ('text_go', txt.get('text_go'),
                 'GO copy — defined but display is currently cleared on flash; add to start_timer() to show it'),
                ('KEY1', cfg.get('KEY1'), "Player 1's reaction key"),
                ('KEY2', cfg.get('KEY2'), "Player 2's reaction key"),
                ('ELAPSED_TIME', cfg.get('ELAPSED_TIME'),
                 'Gap (ms) above which the slower player loses outright; '
                 'within this window the result is randomised'),
                ('RESPONSE_TIMEOUT', cfg.get('RESPONSE_TIMEOUT'),
                 'Seconds to wait for the second press before auto-awarding the win'),
            ],
            'interaction': (
                'Both players press their key. Advances to Screen 10 (forced break) '
                'or Screen 11 (no break) depending on the game condition.'
            ),
        },
        {
            'title': 'Screen 10 — Win with Forced Break',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                "Shown only when the game's break condition is '5 Seconds', '10 Seconds', "
                "or '15 Seconds'. Tells the winner to decide on a blast level during the break. "
                'The screen auto-advances after the break period elapses.'
            ),
            'display': txt.get('text_win_wait', ''),
            'variables': [
                ('text_win_wait', txt.get('text_win_wait'),
                 'Format placeholders: {0}=winner name, {1}=winner name, '
                 '{2}=break seconds, {3}=loser name'),
                ('DEFAULT_GAME_CONDITIONS', cfg.get('DEFAULT_GAME_CONDITIONS'),
                 'Determines FORCED_BREAK_TIME for each game slot (0 / 5 / 10 / 15 seconds)'),
            ],
            'interaction': (
                'No input required — screen advances to Screen 11 (blast level entry) '
                'automatically after FORCED_BREAK_TIME seconds.'
            ),
        },
        {
            'title': 'Screen 11 — Blast Level Entry',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                'The winner selects a blast intensity level from 1 to 9. '
                'The loser is instructed to stand by. A text entry box accepts keyboard input.'
            ),
            'display': txt.get('text_win', ''),
            'variables': [
                ('text_win', txt.get('text_win'),
                 'Format placeholders: {0}=winner name, {1}=winner name (choosing blast), '
                 '{2}=loser name (on standby)'),
            ],
            'interaction': (
                'Winner types a number 1–9 and presses their key. '
                'An invalid entry transitions to Screen 12.'
            ),
        },
        {
            'title': 'Screen 12 — Invalid Blast Level',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                'Shown if the winner enters a character other than 1–9. '
                'The entry field is cleared and the prompt is repeated.'
            ),
            'display': txt.get('text_valid', ''),
            'variables': [
                ('text_valid', txt.get('text_valid'),
                 'Format placeholder: {0}=loser name (still on standby)'),
            ],
            'interaction': 'Winner types a valid number 1–9 and presses their key.',
        },
        {
            'title': 'Screen 13 — Blast Delivery',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                'The selected audio blast plays through the loser\'s headphone channel '
                '(left channel if winner is Player 1, right channel if winner is Player 2). '
                'The screen does not change during playback; the game advances to the next '
                'round automatically once the audio finishes.'
            ),
            'display': (
                '[Audio plays; screen is unchanged during blast]\n\n'
                f'(text_blast = "{txt.get("text_blast", "Standby for blast...")}" is defined '
                f'but not currently shown — add a call to update_text(text_blast) in '
                f'activate_blast() to display it during playback)'
            ),
            'variables': [
                ('text_blast', txt.get('text_blast'),
                 'Standby copy — defined but not currently rendered on screen'),
                ('blast_file', 'radio_static.mp3',
                 'Source audio file; dB levels are adjusted per blast level (1–9) in the '
                 "'Setting up audio' section of the source file"),
            ],
            'interaction': (
                'No input required. Advances to Screen 7 (next round) or '
                'Screen 14 (end of game) automatically.'
            ),
        },
        {
            'title': 'Screen 14 — End-of-Session Summary',
            'bg': f"Black ({cfg.get('BG_COLOUR')})",
            'description': (
                'Shown after all games are complete. Displays a table with wins and '
                'average blast level for each player across each game. '
                'The CSV output file is also saved at this point.'
            ),
            'display': (
                f"{txt.get('text_game_over', 'All games complete!\nThanks for playing!')}\n\n"
                '[Results table]\n'
                '                 | Game 1 (No pause)  | Game 2 (5s pause) | …\n'
                '[Player 1 name]  | X / N wins         | X / N wins        |\n'
                '                 | Avg blast: X.X     | Avg blast: X.X    |\n'
                '[Player 2 name]  | X / N wins         | …                 |\n\n'
                'Press Enter to return to settings'
            ),
            'variables': [
                ('text_game_over', txt.get('text_game_over'),
                 'Title displayed above the results table'),
                ('NUM_ROUNDS', cfg.get('NUM_ROUNDS'),
                 "Denominator in the 'X / N wins' cells"),
            ],
            'interaction': 'Press Enter to return to Screen 1 (Settings) and start a new session.',
        },
    ]

    for screen in screens:
        doc.add_heading(screen['title'], level=2)

        p = doc.add_paragraph()
        bold_label(p, 'Background: ', screen.get('bg', 'Black'))

        doc.add_paragraph(screen['description'])

        p = doc.add_paragraph()
        p.add_run('Display Text:').bold = True
        doc.add_paragraph(screen['display'], style='Quote')

        p = doc.add_paragraph()
        p.add_run('Editable Variables:').bold = True
        for var_name, var_val, desc in screen.get('variables', []):
            bp = doc.add_paragraph(style='List Bullet')
            bp.add_run(f'{var_name}').bold = True
            bp.add_run(f' = {repr(var_val)}  —  {desc}')

        p = doc.add_paragraph()
        bold_label(p, 'Player Interaction: ', screen.get('interaction', ''))

        p = doc.add_paragraph()
        p.add_run('Screenshot:').bold = True
        add_screenshot_placeholder(doc)

    doc.save(OUTPUT_FILE)
    print(f'Saved: {OUTPUT_FILE}')


if __name__ == '__main__':
    main()
